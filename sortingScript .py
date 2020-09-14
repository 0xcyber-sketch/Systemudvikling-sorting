#!/usr/bin/env python
# -*- coding: UTF-8 -*-

import docx
import os
import re

# Stor path withouth the file
pathWF = 'C:\\Users\\Emil\\Desktop\\1. Semester\\SU1'

# Takes input name for file
path = input("Name for old word file: ")

if __name__ == '__main__':
    # Bool for with or wothout .docx
    withOutDocx = None

    if ".docx" not in path:
        # Stores length of file name for later
        fileLength = len(path)
        # Joins the path
        path = os.path.join(pathWF, path)
        path = path + ".docx"
        withOutDocx = True
    else:
        fileLength = len(path)
        path = os.path.join(pathWF, path)
        withOutDocx = False

    # Opens the old/original doc
    doc = docx.Document(path)

    # start value for enumaration
    x = 0

    # Lists of bullet points
    list_bullet = []

    # New list for splitted questions
    newList = []

    # String to find amount of bullet points
    string = ""

    # Empty string for storing new string
    newString = ""

    # Enumerate through the document
    for i in enumerate(doc.paragraphs, x):
        # Look for a long paragraph
        if len(doc.paragraphs[i[0]].text) > 13:
            # Make a string to look for Bullets
            string = doc.paragraphs[i[0]].text
            # Check for x amount of bullet in string and append to my bullets list
            for m in re.finditer("•", string):
                k = m.start(), m.end()
                list_bullet.append(k)
        x += 1

    # Get the lenght of the bullet list
    lenght_of_list = len(list_bullet)

    # Make every question a element in a list
    for i in range(lenght_of_list):
        newList = string.split("•")

    # Reverse the list to get the proper order
    newList = newList[::-1]

     # Add each elemt in the list to on string
    for i in range(lenght_of_list):
        newString = "•" + newList[i] + "\n" + newString

    # Make a new docs
    d = docx.Document()

    # Change style
    style = d.styles['Normal']
    font = style.font
    font.name = 'Arial'

    # Add the string to the docs
    d.add_paragraph(newString)

    # Prompt user for name for new doc file
    name = input("New name for the file: ")

    # Makes the new path for the file. Checks if it needs to consider the .docx
    if withOutDocx == True:
        new_Path = path[:-(fileLength + 5)]

    if withOutDocx == False:
        new_Path = path[:-fileLength]

    new_Path = os.path.join(new_Path, name)
    new_Path = new_Path + ".docx"

    # Saves the new document at the new path
    d.save(new_Path)
